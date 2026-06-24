"""Shared pytest fixtures."""
import io

import docx
import fitz
import pytest
from fastapi.testclient import TestClient

from app.main import app


@pytest.fixture(scope="session")
def client() -> TestClient:
    return TestClient(app)


@pytest.fixture
def sample_docx_bytes() -> bytes:
    d = docx.Document()
    d.add_heading("Jane Developer", 0)
    d.add_paragraph("jane@example.com")
    d.add_heading("Experience", 1)
    d.add_paragraph("Led a team and reduced costs by 40% using Go and C#.")
    d.add_heading("Skills", 1)
    d.add_paragraph("Python, Go, C#, ML, AWS")
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


@pytest.fixture
def sample_pdf_bytes() -> bytes:
    p = fitz.open()
    page = p.new_page()
    page.insert_text((72, 72), "Resume. Managed AWS infrastructure. Skills: Java, SQL, ML.")
    data = p.tobytes()
    p.close()
    return data
