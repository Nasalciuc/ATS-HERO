"""Tests for the upgraded content findings + the three new detectors."""
import io

import docx
import fitz

from app.analyzers.ats_scorer import score_text
from app.analyzers.format_analyzer import analyze_format
from app.analyzers.job_matcher import match_job
from app.analyzers.keyword_extractor import extract_skills
from app.analyzers.nlp_engine import get_engine
from app.analyzers.region_checker import check_region
from app.models.requests import RegionSignals


# ---------- content findings (weak phrases / clichés / extended metrics) ----------

def test_weak_phrase_findings():
    cv = "Responsible for APIs. Helped with the backend. A passionate team player."
    r = score_text(cv, get_engine())
    types = {f.type for f in r.findings}
    assert "weak_phrase" in types
    assert "cliche" in types
    weak_texts = {f.text for f in r.findings if f.type == "weak_phrase"}
    assert "responsible for" in weak_texts
    assert "helped with" in weak_texts
    cliche_texts = {f.text for f in r.findings if f.type == "cliche"}
    assert "team player" in cliche_texts or "passionate" in cliche_texts


def test_findings_have_line_numbers_and_hints():
    cv = "Summary\nResponsible for APIs."
    r = score_text(cv, get_engine())
    weak = [f for f in r.findings if f.type == "weak_phrase"]
    assert weak and weak[0].line == 2
    assert weak[0].suggestion and weak[0].why


def test_extended_metrics_detected():
    # engineering metrics that the old regex missed
    cv = "Built a gRPC API serving 40k req/s at p99 180ms. Cut page load 4.2s to 1.1s."
    r = score_text(cv, get_engine())
    impact = next(s for s in r.signals if s.key == "quantified_impact")
    assert impact.score > 0


def test_clean_cv_has_no_weak_findings():
    cv = "Led a team of 5 and reduced latency by 40% using Go."
    r = score_text(cv, get_engine())
    assert not any(f.type == "weak_phrase" for f in r.findings)


# ---------- job-fit ----------

def test_jobfit_basic_overlap():
    cv = "Experienced with Python, Go, and Kubernetes on AWS."
    jd = "We need Go, Kubernetes, AWS, and gRPC."
    r = match_job(cv, jd)
    assert "Go" in r.matched and "AWS" in r.matched
    assert "gRPC" in r.missing
    assert 0 < r.match_score < 100


def test_jobfit_acronym_equivalence():
    cv = "Strong background in machine learning and continuous integration."
    jd = "Looking for ML and CI experience."
    r = match_job(cv, jd)
    # ML matched via 'machine learning', CI via 'continuous integration'
    assert "ML" in r.matched
    assert "ML" not in r.missing


def test_jobfit_perfect_match():
    cv = "Go, Kubernetes, AWS."
    jd = "Go, Kubernetes, AWS."
    r = match_job(cv, jd)
    assert r.match_score == 100.0
    assert r.missing == []


# ---------- format / parseability ----------

def _docx_with_table() -> bytes:
    d = docx.Document()
    d.add_paragraph("Jane Developer")
    t = d.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "Python"
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def _clean_docx() -> bytes:
    d = docx.Document()
    d.add_heading("Jane Developer", 0)
    d.add_paragraph("Led a team and shipped features.")
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def test_format_flags_tables():
    r = analyze_format(_docx_with_table(), "cv.docx", None)
    assert r.format == "docx"
    assert any(w.type == "tables" for w in r.warnings)
    assert r.ats_safe is False   # tables are high severity


def test_format_clean_docx_is_safe():
    r = analyze_format(_clean_docx(), "cv.docx", None)
    assert r.ats_safe is True
    assert r.page_count >= 1


def test_format_plain_pdf_is_safe():
    p = fitz.open()
    page = p.new_page()
    page.insert_text((72, 72), "Jane Developer. Led a team and reduced costs by 40%.")
    data = p.tobytes()
    p.close()
    r = analyze_format(data, "cv.pdf", None)
    assert r.format == "pdf"
    assert r.ats_safe is True


# ---------- region (global) ----------

def test_region_us_flags_photo():
    r = check_region("US", RegionSignals(has_photo=True))
    assert any(d.field == "photo" and d.severity == "high" for d in r.deviations)


def test_region_germany_flags_missing_photo():
    r = check_region("DE", RegionSignals(has_photo=False))
    assert any(d.field == "photo" for d in r.deviations)


def test_region_silent_when_compliant():
    # US baseline: no photo, no personal data, 1 page -> no deviations
    r = check_region("US", RegionSignals(has_photo=False, has_personal_data=False, page_count=1))
    assert r.deviations == []


def test_region_unknown_signals_not_flagged():
    # everything None -> nothing to flag, even for a strict locale
    r = check_region("DE", RegionSignals())
    assert r.deviations == []


def test_region_length_uses_target_norm():
    # South Africa allows 3 pages; 3 pages should NOT be flagged there
    assert check_region("ZA", RegionSignals(page_count=3)).deviations == [] or all(
        d.field != "length" for d in check_region("ZA", RegionSignals(page_count=3)).deviations
    )
    # but 3 pages IS flagged for the US (2-page norm)
    assert any(d.field == "length" for d in check_region("US", RegionSignals(page_count=3)).deviations)


def test_region_text_derives_gdpr_for_italy():
    cv = "Work experience...\nI authorize the processing of my personal data per GDPR."
    r = check_region("IT", RegionSignals(), cv_text=cv)
    # GDPR clause present -> should NOT be flagged for Italy
    assert all(d.field != "gdpr_clause" for d in r.deviations)


# ---------- multi-industry keyword coverage (article-driven update) ----------

def test_healthcare_skills_detected():
    cv = "Nursing professional skilled in case management, healthcare, pharmacy, and CPR."
    low = " ".join(extract_skills(cv)).lower()
    assert "nursing" in low and "case management" in low and "cpr" in low


def test_finance_skills_detected():
    cv = "Prepared financial statements, owned accounts payable, ensured GAAP compliance."
    low = " ".join(extract_skills(cv)).lower()
    assert "financial statements" in low and "accounts payable" in low and "gaap" in low


def test_marketing_skills_detected():
    cv = "Ran digital marketing and social media campaigns; tracked key performance indicators."
    low = " ".join(extract_skills(cv)).lower()
    assert "digital marketing" in low and "social media" in low


def test_jobfit_industry_acronym():
    cv = "Built dashboards tracking key performance indicators and managed CRM data."
    jd = "Need KPI reporting and CRM experience."
    r = match_job(cv, jd)
    assert "KPI" in r.matched   # via the spelled-out phrase
    assert "CRM" in r.matched   # direct match
