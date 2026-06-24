"""Format / parseability analysis.

Inspects the FILE STRUCTURE (not the extracted text), because the things that break ATS
parsers — columns, tables, images, header/footer content, image-only PDFs — are invisible
once text is flattened. PDF via PyMuPDF, DOCX via python-docx. High-confidence signals are
reported as-is; the multi-column heuristic is flagged as 'probable' to avoid false positives.
"""
import io

import fitz  # PyMuPDF

from ..models.responses import Finding, FormatResponse


def analyze_format(data: bytes, filename: str, content_type: str | None) -> FormatResponse:
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    ct = (content_type or "").lower()
    if ext == "pdf" or ct == "application/pdf":
        return _analyze_pdf(data)
    if ext == "docx" or "word" in ct or "wordprocessing" in ct:
        return _analyze_docx(data)
    raise ValueError("Unsupported file type — use DOCX or PDF")


def _page_count(data: bytes, filetype: str) -> int:
    try:
        doc = fitz.open(stream=data, filetype=filetype)
        n = doc.page_count
        doc.close()
        return n
    except Exception:
        return 0


def _looks_multicolumn(page) -> bool:
    blocks = [b for b in page.get_text("blocks") if b[6] == 0 and b[4].strip()]
    if len(blocks) < 6:
        return False
    mid = page.rect.width / 2
    left = [b for b in blocks if b[2] < mid]
    right = [b for b in blocks if b[0] > mid]
    return len(left) >= 3 and len(right) >= 3


def _analyze_pdf(data: bytes) -> FormatResponse:
    doc = fitz.open(stream=data, filetype="pdf")
    warnings: list[Finding] = []
    try:
        pages = doc.page_count
        total_text = 0
        total_images = 0
        fonts: set[str] = set()
        multicolumn = False
        for page in doc:
            total_text += len(page.get_text().strip())
            total_images += len(page.get_images(full=True))
            for f in page.get_fonts(full=True):
                fonts.add(f[3])
            if _looks_multicolumn(page):
                multicolumn = True
    finally:
        doc.close()

    if total_images > 0 and total_text < 50:
        warnings.append(Finding(
            type="image_only_pdf", severity="high", text="image-based PDF",
            why="The file looks like a scanned/image PDF with no extractable text — ATS cannot read it.",
            suggestion="Export a text-based PDF from Word or Google Docs (real selectable text).",
        ))
    elif total_images > 0:
        warnings.append(Finding(
            type="images", severity="medium", text=f"{total_images} image(s)",
            why="ATS cannot extract text from images, icons, skill bars, or logos.",
            suggestion="Remove graphics; put skills and contact info as real text.",
        ))
    if multicolumn:
        warnings.append(Finding(
            type="multi_column", severity="high", text="probable multi-column layout",
            why="Probable two-column layout — parsers flatten columns and can read content out of order.",
            suggestion="Switch to a single-column layout.",
        ))
    if pages > 2:
        warnings.append(Finding(
            type="length", severity="low", text=f"{pages} pages",
            why="More than two pages; many roles and markets expect 1-2.",
            suggestion="Trim to the most relevant 1-2 pages (varies by region — see region check).",
        ))
    if len(fonts) > 4:
        warnings.append(Finding(
            type="fonts", severity="low", text=f"{len(fonts)} fonts",
            why="Many fonts can look noisy and occasionally trip older parsers.",
            suggestion="Use one or two standard fonts (Arial, Calibri, Times New Roman).",
        ))

    ats_safe = not any(w.severity == "high" for w in warnings)
    return FormatResponse(
        format="pdf", ats_safe=ats_safe, page_count=pages, warnings=warnings,
        details={"text_chars": total_text, "images": total_images, "fonts": len(fonts), "multicolumn": multicolumn},
    )


def _analyze_docx(data: bytes) -> FormatResponse:
    import docx  # python-docx

    document = docx.Document(io.BytesIO(data))
    warnings: list[Finding] = []

    n_tables = len(document.tables)
    n_images = len(document.inline_shapes)

    header_text = ""
    for section in document.sections:
        try:
            header_text += " ".join(p.text for p in section.header.paragraphs)
            header_text += " ".join(p.text for p in section.footer.paragraphs)
        except Exception:
            pass
    header_has_content = bool(header_text.strip())

    # Text boxes live in the document XML as w:txbxContent.
    has_textbox = b"txbxContent" in data

    if n_tables > 0:
        warnings.append(Finding(
            type="tables", severity="high", text=f"{n_tables} table(s)",
            why="Many ATS (Workday especially) cannot read content inside tables — even a Skills table.",
            suggestion="Remove tables; lay out content as plain paragraphs and lists.",
        ))
    if header_has_content:
        warnings.append(Finding(
            type="header_footer", severity="high", text="content in header/footer",
            why="Many ATS ignore headers/footers — contact info placed there can be lost entirely.",
            suggestion="Move your name and contact details into the document body.",
        ))
    if has_textbox:
        warnings.append(Finding(
            type="text_box", severity="high", text="text box(es)",
            why="Text inside text boxes is frequently skipped by ATS parsers.",
            suggestion="Replace text boxes with normal paragraphs.",
        ))
    if n_images > 0:
        warnings.append(Finding(
            type="images", severity="medium", text=f"{n_images} image(s)",
            why="ATS cannot extract text from images, icons, or logos.",
            suggestion="Remove graphics; keep skills and contact info as real text.",
        ))

    pages = _page_count(data, "docx")
    if pages > 2:
        warnings.append(Finding(
            type="length", severity="low", text=f"{pages} pages",
            why="More than two pages; many roles and markets expect 1-2.",
            suggestion="Trim to the most relevant 1-2 pages (varies by region — see region check).",
        ))

    ats_safe = not any(w.severity == "high" for w in warnings)
    return FormatResponse(
        format="docx", ats_safe=ats_safe, page_count=pages, warnings=warnings,
        details={"tables": n_tables, "images": n_images, "header_footer": header_has_content, "text_box": has_textbox},
    )
